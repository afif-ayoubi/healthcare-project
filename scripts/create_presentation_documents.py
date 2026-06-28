from __future__ import annotations

from pathlib import Path
import sys

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT=Path(__file__).resolve().parents[1]
sys.path.insert(0,str(ROOT/'scripts'))
from create_documents import (
    setup_document, add_cover, add_page_break, add_kicker, add_heading, add_body,
    add_bullets, add_callout, add_table, set_run_font, NAVY, TEAL, MINT, CORAL,
    GOLD, SLATE, PALE, LIGHT, BORDER
)

PRES=ROOT/'presentation'
PRES.mkdir(parents=True,exist_ok=True)

speaker_notes = [
("Slide 1 · Opening", "0:00–0:40", "Tuberculosis is preventable and curable, yet it remains one of the world’s largest infectious-disease burdens. My project turns WHO-derived country-level estimates into a Streamlit decision-support dashboard. In the reproducible 2024 snapshot, about 10.55 million people developed TB, 8.33 million new and relapse cases were notified, and calculated diagnosis-and-treatment coverage was 78.9%. The purpose is to show not only where TB is high, but where the health-system response may still be incomplete."),
("Slide 2 · Problem and objective", "0:40–1:25", "The analytical problem is the gap between estimated disease burden and the people who appear in routine services. I calculate the notification difference as estimated incident cases minus notified new and relapse cases. It is deliberately labelled as an estimate and a surveillance signal, because it does not prove that every person in the difference was untreated. The dashboard answers four practical questions: where burden is highest, how trends are changing, who and which subtypes are most affected, and where a healthcare system should investigate first."),
("Slide 3 · Data and method", "1:25–2:10", "The workflow is reproducible from raw input to dashboard. I retained the raw snapshots, standardized identifiers and numeric fields in pandas, merged the indicators into one country-year model, and recalculated global and regional rates from the country data. Automated checks confirm unique country-year records, the 2000-to-2024 range, non-negative burden values and male/female age-sex data. The final package contains the code, raw and processed files, a data dictionary, source hashes, and validation output."),
("Slide 4 · Global trend", "2:10–3:00", "The core time series shows substantial improvement. Notifications more than doubled since 2000, and calculated coverage rose from 32.4% to 78.9%. However, the 2024 snapshot still contains a 2.22 million difference between estimated burden and notified cases. The 2020 disruption is especially important: notifications fell by 18.1% from 2019 while estimated incidence moved much less. This shows that service access and reporting can deteriorate rapidly even when the underlying disease burden does not disappear."),
("Slide 5 · Geography and gaps", "3:00–3:50", "Geography changes the interpretation. Indonesia had the largest absolute notification difference, about 241 thousand, while India had the largest total burden but calculated coverage above 92%. Myanmar illustrates a different risk: a smaller absolute burden, but coverage below 44%. At regional level, the Western Pacific had the largest absolute notification difference. The dashboard therefore avoids ranking countries on one indicator alone and encourages users to combine absolute volume, population rate and service coverage."),
("Slide 6 · Demographics and subtypes", "3:50–4:35", "The age-sex analysis adds a population lens. Estimated male incidence was about one and a half times female incidence, with the largest adult burden in ages 25 to 54. The subtype analysis adds a clinical-programme lens. HIV-associated TB was concentrated in sub-Saharan Africa, especially South Africa, while drug-resistant TB followed a different pattern led by India and several European and Asian settings. These are not interchangeable problems: TB–HIV requires integrated testing and treatment pathways, while resistance requires rapid susceptibility testing and effective regimens."),
("Slide 7 · Recommendations and close", "4:35–5:00", "The dashboard leads to four actions: prioritize high-volume detection gaps; protect TB services during disruption; integrate TB–HIV and drug-resistance pathways according to the local pattern; and tailor outreach while governing the data carefully. The value of the project is the sequence: find where burden and coverage gaps are concentrated, target the population and programme response, and act with measurable interventions. The app, data, code, manual, report and presentation are ready for deployment; the final step is publishing through my Streamlit account."),
]


def build_script():
    doc=setup_document("Global TB Project - Five-Minute Presentation Script","Five-minute pitch script")
    add_cover(doc,"Global Tuberculosis Burden & Treatment Gaps","Five-minute pitch + three-minute Q&A preparation","Presentation script",include_summary=True)
    add_page_break(doc)
    add_kicker(doc,"Pitch map")
    add_heading(doc,"Five-minute delivery plan",1)
    add_body(doc,"Target speaking time: approximately 5:00. The script is designed for a calm pace, with one main message per slide. The presentation should sound conversational rather than memorized.")
    rows=[]
    for i,(title,timing,text) in enumerate(speaker_notes,1):
        rows.append([str(i),title.replace(f"Slide {i} · ",""),timing,f"{len(text.split())} words"])
    add_table(doc,["Slide","Purpose","Target time","Length"],rows,widths=[.55,3.45,1.15,1.15],font_size=8.0)
    add_heading(doc,"Three messages to repeat",2)
    add_bullets(doc,[
        "The dashboard combines burden, rates and coverage instead of ranking countries on one measure.",
        "The notification difference is a surveillance signal, not a verified count of untreated people.",
        "The final recommendations depend on geography, demographics, TB–HIV co-burden and drug resistance.",
    ],size=9,space_after=2)
    add_callout(doc,"Delivery cue","Pause after each headline number. Point to the visual before explaining it. Do not read chart labels line by line.",fill="EAF5F7",accent=MINT,font_size=9)

    # Slides 1-3
    add_page_break(doc)
    add_kicker(doc,"Script · Part 1")
    add_heading(doc,"Opening, problem and method",1)
    for title,timing,text in speaker_notes[:3]:
        add_heading(doc,f"{title}  |  {timing}",2)
        add_body(doc,text)
        if title.startswith("Slide 1"):
            add_callout(doc,"Emphasis","Say “reproducible 2024 snapshot” before the numbers so the audience understands these are the dashboard aggregates.",fill="FFF6ED",accent=CORAL,font_size=8.2)
        elif title.startswith("Slide 2"):
            add_callout(doc,"Do not overclaim","Avoid saying that all 2.22 million people were untreated. Say “estimated notification difference” or “coverage signal.”",fill="FFF6ED",accent=CORAL,font_size=8.2)

    # Slides 4-5
    add_page_break(doc)
    add_kicker(doc,"Script · Part 2")
    add_heading(doc,"Global trend and geographic priorities",1)
    for title,timing,text in speaker_notes[3:5]:
        add_heading(doc,f"{title}  |  {timing}",2)
        add_body(doc,text)
    add_heading(doc,"Transition sentence",2)
    add_callout(doc,"Use this line","“The global trend tells us whether progress is occurring; the geography tells us where the next unit of effort may have the greatest value.”",fill="EAF5F7",accent=TEAL,font_size=9)
    add_heading(doc,"Numbers to remember",2)
    add_table(doc,["Finding","Value"],[
        ["2024 estimated incident cases","10.55 million"],
        ["2024 notified new and relapse cases","8.33 million"],
        ["2024 calculated coverage","78.9%"],
        ["2024 estimated notification difference","2.22 million"],
        ["2019→2020 notification change","−18.1%"],
        ["Largest country-level difference","Indonesia · about 241,000"],
    ],widths=[4.7,2.0],font_size=8.3)

    # Slides 6-7
    add_page_break(doc)
    add_kicker(doc,"Script · Part 3")
    add_heading(doc,"Demographics, recommendations and close",1)
    for title,timing,text in speaker_notes[5:]:
        add_heading(doc,f"{title}  |  {timing}",2)
        add_body(doc,text)
    add_heading(doc,"Final sentence",2)
    add_callout(doc,"Close strongly","“This dashboard helps a healthcare decision-maker move from finding the problem, to targeting the response, to acting with measurable evidence.”",fill="EAF5F7",accent=MINT,font_size=9)
    add_heading(doc,"Presentation-day checklist",2)
    add_bullets(doc,[
        "Open the published dashboard and pitch deck before class; keep the dashboard on the executive overview.",
        "Replace [Student Name] and paste the final Streamlit URL in the deck, report and README.",
        "Practice once at normal pace and once with a strict five-minute timer.",
        "During Q&A, answer the question first, then provide one supporting number or limitation.",
    ],size=8.5,space_after=1)
    out=PRES/'TB_Presentation_Script.docx'; doc.save(out); return out


qa_groups=[
("Concept and interpretation",[
("What exactly is the “treatment gap”?","The dashboard’s core measurable gap is diagnosis-and-treatment coverage: notified new and relapse cases divided by estimated incident cases. The absolute difference is called the estimated notification difference. I avoid claiming that every person in the difference was untreated because incidence is estimated and notifications are programme counts."),
("Why is TB related to HIV?","HIV weakens immunity, increasing the chance that TB infection progresses to active disease. They are different diseases, so TB–HIV is analysed as a co-epidemic requiring integrated testing and care."),
("What is RR/MDR-TB?","RR-TB is resistant to rifampicin, a key TB medicine. MDR-TB is resistant to at least rifampicin and isoniazid. These cases require rapid resistance testing and different treatment pathways."),
("Why use both counts and rates?","Counts identify where the largest service volume is needed. Rates standardize for population size and identify intense risk. Coverage adds the health-system reach dimension."),
]),
("Data and methodology",[
("Where did the data come from?","The project uses WHO Global Tuberculosis Programme indicators. The package includes raw WHO-derived snapshots, an official WHO report/repository cross-check, source URLs, retrieval dates and SHA-256 hashes."),
("Why did you use a mirror rather than the direct WHO CSV?","The WHO data page is the authoritative source, but its direct CSV endpoint was not retrievable in the development runtime. I used a public WHO-derived DDF mirror, retained full provenance, and validated results against the current WHO report and official 2025 repository."),
("Why are your global totals 10.55M and 1.20M, while WHO says 10.7M and 1.23M?","WHO publishes rounded global estimates with uncertainty intervals and updates country time series. My figures are reproducible sums of the country-level snapshot. The dashboard discloses this difference instead of presenting the two as identical."),
("What data cleaning did you perform?","I standardized country identifiers and names, converted years and indicators to numeric form, merged 15 indicator files, recalculated rates and coverage, selected non-overlapping age bands, and created country, regional and global tables."),
("How did you validate the data?","Automated checks tested unique country-year records, year range, non-negative burden values, ISO-code completeness, expected WHO regions, and the presence of male and female age-sex estimates."),
]),
("Analysis, prediction and limitations",[
("Why can calculated coverage exceed 100%?","Incidence is estimated and notifications are reported. Different uncertainty, timing, revisions or cross-border diagnosis can produce a numerator above the estimated denominator. The dashboard preserves the value but uses a non-negative gap field for ranking."),
("Why use a simple linear forecast?","Prediction is optional bonus work. I chose an interpretable baseline, report holdout MAE, compare it with a last-value baseline, and label the result as an exploratory extrapolation rather than a policy or clinical forecast."),
("Why is age/sex only for 2024?","That is the available cross-sectional age-sex estimate in the integrated project package. I state the limitation and do not imply a demographic time trend."),
("Why not include treatment success outcomes?","The current version focuses on burden, notifications and diagnosis-and-treatment coverage. Treatment cohorts use different reference years and require careful alignment. They are a logical extension, but omitting them keeps this version methodologically coherent."),
("Can the dashboard prove causes?","No. It uses ecological country-level data and supports prioritization, not causal inference or individual risk prediction."),
("What is the single most important finding?","Progress is visible, but the response is uneven: a global 78.9% calculated coverage still leaves a large notification difference, and priority changes depending on burden volume, rate, HIV co-burden and drug resistance."),
("What would you improve next?","I would add treatment-cohort outcomes, uncertainty intervals, subnational data for selected countries, and stronger forecast models only after verifying sufficient and comparable time-series data."),
]),
]


def build_qa():
    doc=setup_document("Global TB Project - Q&A Notes","Three-minute Q&A preparation")
    add_cover(doc,"Global Tuberculosis Burden & Treatment Gaps","Likely questions, defensible answers and wording guardrails","Q&A preparation",include_summary=True)
    add_page_break(doc)
    add_kicker(doc,"Q&A strategy")
    add_heading(doc,"Answer in three steps",1)
    add_table(doc,["Step","What to do","Example"],[
        ["1. Answer","Give the direct answer in one sentence.","“No, the dashboard does not prove causality.”"],
        ["2. Evidence","Support it with one method or number.","“It uses ecological country-level data.”"],
        ["3. Boundary","State the limitation or next step.","“A causal study would need individual or quasi-experimental data.”"],
    ],widths=[1.0,2.7,3.0],font_size=8.2)
    add_heading(doc,"Wording guardrails",2)
    add_bullets(doc,[
        "Say “estimated incident cases,” not “confirmed cases.”",
        "Say “notified new and relapse cases,” not simply “treated patients.”",
        "Say “estimated notification difference” or “coverage signal,” not “all untreated people.”",
        "Say “exploratory projection,” not “prediction of what will happen.”",
        "Say “associated geographic pattern,” not “cause.”",
    ],size=9,space_after=2)
    add_callout(doc,"Best fallback","“The dashboard is designed to prioritize investigation. It does not replace programme validation or clinical judgment.”",fill="FFF6ED",accent=CORAL,font_size=9)

    # Group 1 and part group2
    add_page_break(doc)
    add_kicker(doc,"Likely questions · Part 1")
    for group,qs in qa_groups[:2]:
        add_heading(doc,group,1)
        for q,a in qs:
            add_heading(doc,q,2)
            add_body(doc,a)

    # Remaining group
    add_page_break(doc)
    add_kicker(doc,"Likely questions · Part 2")
    group,qs=qa_groups[2]
    add_heading(doc,group,1)
    for q,a in qs:
        add_heading(doc,q,2)
        add_body(doc,a)
    add_heading(doc,"Source references to mention",2)
    add_bullets(doc,[
        "WHO Global tuberculosis report 2025 and its top findings.",
        "WHO Tuberculosis data page and data dictionary.",
        "GTB-TME official 2025 report GitHub repository.",
        "Project data manifest, preparation script and validation JSON.",
    ],size=8.5,space_after=1)
    out=PRES/'TB_QA_Notes.docx'; doc.save(out); return out

if __name__=='__main__':
    print(build_script())
    print(build_qa())
