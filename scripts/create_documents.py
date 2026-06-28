from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(parents=True, exist_ok=True)
CH = ROOT / "outputs" / "charts"
SCREEN = ROOT / "outputs" / "screenshots"
P = ROOT / "data" / "processed"

NAVY = "0B1F33"
TEAL = "0E7490"
MINT = "14B8A6"
CORAL = "E76F51"
GOLD = "F4A261"
SLATE = "64748B"
PALE = "EEF6F8"
LIGHT = "F7FAFC"
BORDER = "DDE7EC"
WHITE = "FFFFFF"
FONT = "Liberation Sans"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    """Set borders: top/bottom/start/end={'val':'single','sz':'6','color':'DDE7EC'}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:" + edge
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:" + key), str(edge_data[key]))


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_col_widths(table, widths: list[float]):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = FONT
    run.font.size = Pt(8)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 15, TEAL, 0, 10),
        ("Heading 1", 20, NAVY, 4, 8),
        ("Heading 2", 13, TEAL, 4, 5),
        ("Heading 3", 10.5, NAVY, 3, 3),
    ]:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Kicker" not in styles:
        st = styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(8.2)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(TEAL)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True
    if "CaptionCustom" not in styles:
        st = styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(7.7)
        st.font.italic = True
        st.font.color.rgb = RGBColor.from_string(SLATE)
        st.paragraph_format.space_before = Pt(2)
        st.paragraph_format.space_after = Pt(5)
    if "Small" not in styles:
        st = styles.add_style("Small", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(8)
        st.font.color.rgb = RGBColor.from_string(SLATE)
        st.paragraph_format.space_after = Pt(3)
        st.paragraph_format.line_spacing = 1.0

    # Header/footer
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(subtitle.upper())
    set_run_font(r, 7.5, True, TEAL)
    footer = section.footer
    ft = footer.paragraphs[0]
    ft.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = ft.add_run("MSBA382 · Healthcare Analytics  |  Global TB Project")
    set_run_font(r, 7.5, False, SLATE)
    add_page_number(ft)

    core = doc.core_properties
    core.title = title
    core.subject = subtitle
    core.author = "[Student Name]"
    core.comments = "Prepared for MSBA382 Healthcare Analytics"
    return doc


def add_kicker(doc, text):
    p = doc.add_paragraph(style="Kicker")
    p.add_run(text.upper())
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text, bold_prefix: str | None = None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items: Iterable[str], level=0, size=9.2, space_after=2):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.left_indent = Inches(0.18 + level * 0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(item)
        set_run_font(r, size=size, color=NAVY)


def add_callout(doc, title: str, body: str, fill=PALE, accent=TEAL, font_size=8.8):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    set_cell_border(cell, start={"val": "single", "sz": "18", "color": accent})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title + "  ")
    set_run_font(r, font_size, True, accent)
    r = p.add_run(body)
    set_run_font(r, font_size, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers: list[str], rows: list[list], widths=None, font_size=8.1, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        set_cell_shading(c, header_fill)
        set_cell_margins(c, top=75, bottom=75, start=70, end=70)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        set_run_font(r, font_size, True, WHITE)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            c = cells[j]
            set_cell_shading(c, WHITE if i % 2 == 0 else LIGHT)
            set_cell_margins(c, top=65, bottom=65, start=70, end=70)
            set_cell_border(c, bottom={"val": "single", "sz": "4", "color": BORDER})
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, font_size, False, NAVY)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        set_col_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_picture(doc, path: Path, width: float, caption: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph(caption, style="CaptionCustom")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_cover(doc: Document, title: str, subtitle: str, label: str, include_summary=True):
    # Hide header on first page
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("MSBA382 · HEALTHCARE ANALYTICS")
    set_run_font(r, 9, True, TEAL)

    band = doc.add_table(rows=1, cols=1)
    cell = band.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=360, bottom=360, start=260, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, 28, True, WHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(subtitle)
    set_run_font(r, 14, False, "CDEDF0")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label.upper())
    set_run_font(r, 10, True, CORAL)

    if include_summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(
            "A decision-support tool for understanding where tuberculosis burden is concentrated, "
            "how trends have changed, and where diagnosis-and-treatment coverage remains incomplete."
        )
        set_run_font(r, 12.2, False, NAVY)

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.LEFT
    labels = ["Name", "Major", "Department", "Date"]
    values = ["[Student Name]", "Master of Science in Business Analytics", "Suliman S. Olayan School of Business", "23 June 2026"]
    for i, (a, b) in enumerate(zip(labels, values)):
        for c in info.rows[i].cells:
            set_cell_margins(c, top=90, bottom=90, start=90, end=90)
            set_cell_border(c, bottom={"val": "single", "sz": "5", "color": BORDER})
        set_cell_shading(info.cell(i, 0), PALE)
        set_cell_shading(info.cell(i, 1), WHITE)
        r = info.cell(i, 0).paragraphs[0].add_run(a)
        set_run_font(r, 9, True, SLATE)
        r = info.cell(i, 1).paragraphs[0].add_run(b)
        set_run_font(r, 9.2, False, NAVY)
    set_col_widths(info, [1.35, 5.5])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_callout(
        doc,
        "Project snapshot",
        "Analytical period 2000–2024 · 217 countries/territories · six WHO regions · age/sex cross-section for 2024 · reproducible source manifest and validation checks.",
        fill="EAF5F7",
        accent=MINT,
        font_size=9,
    )
    doc.add_paragraph().paragraph_format.space_before = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared as an individual master’s-level healthcare analytics project")
    set_run_font(r, 8.5, False, SLATE, italic=True)


# ----------------------- MANUAL -----------------------
def build_manual():
    doc = setup_document(
        "Global Tuberculosis Burden and Treatment Gaps - Dashboard Manual",
        "Consultant-to-client dashboard manual",
    )
    add_cover(
        doc,
        "Global Tuberculosis Burden & Treatment Gaps",
        "Trends, demographic disparities and geographic patterns",
        "Consultant-to-client dashboard manual",
    )

    # Page 2
    add_page_break(doc)
    add_kicker(doc, "01 · Data foundation")
    add_heading(doc, "Data sources and analytical definitions", 1)
    add_body(doc, "The dashboard uses secondary, country-level public-health data. Raw snapshots are preserved, and every processed file can be recreated with the included Python scripts.")

    add_heading(doc, "Source package", 2)
    source_rows = [
        ["WHO TB burden indicators", "Incidence, mortality, population, HIV-associated TB and RR/MDR-TB", "2000–2024"],
        ["WHO case notifications", "Notified new and relapse TB cases; calculated coverage", "2000–2024"],
        ["WHO age-sex estimates", "Estimated incident TB by age band and sex", "2024"],
        ["WHO region lookup", "Maps ISO3 country codes to six WHO regions", "Current snapshot"],
    ]
    add_table(doc, ["Dataset", "Role in the dashboard", "Coverage"], source_rows, widths=[1.65, 3.9, 1.15], font_size=7.8)

    add_heading(doc, "Core measures", 2)
    measure_rows = [
        ["Estimated incident TB", "People estimated to have developed TB during the year", "Count / rate per 100,000"],
        ["Notified new & relapse", "Cases officially reported to national TB programmes", "Count / rate per 100,000"],
        ["Calculated coverage", "Notified new/relapse cases ÷ estimated incident cases × 100", "%"],
        ["Notification difference", "Estimated incident cases − notified new/relapse cases", "Count"],
        ["Estimated TB mortality", "Estimated deaths from TB", "Count / rate per 100,000"],
        ["HIV-associated TB", "Incident TB among people living with HIV", "Count / share"],
        ["RR/MDR-TB", "Rifampicin-resistant or multidrug-resistant incident TB", "Count / share"],
    ]
    add_table(doc, ["Measure", "Meaning", "Unit"], measure_rows, widths=[1.7, 4.1, 1.1], font_size=7.5, header_fill=TEAL)

    add_callout(
        doc,
        "Interpretation guardrail",
        "The notification difference is a surveillance and service-coverage signal. It must not be presented as proof that every person in the difference was untreated. Burden values are estimates with uncertainty, while notifications are reported programme counts.",
        fill="FFF6ED",
        accent=CORAL,
        font_size=8.2,
    )
    add_body(doc, "Reproducibility: data/processed/data_manifest.csv records the source URL, retrieval date, dimensions, year range and SHA-256 hash for every raw input.", style="Small")

    # Page 3
    add_page_break(doc)
    add_kicker(doc, "02 · Tool navigation")
    add_heading(doc, "How to use the Streamlit dashboard", 1)
    add_picture(doc, SCREEN / "dashboard_preview.png", 6.9, "Figure 1. Static preview of the dashboard’s executive overview. The published app is interactive.")
    tab_rows = [
        ["Executive overview", "KPI cards, burden vs notifications, rate trends and prioritized country table"],
        ["Global map", "Choropleth for incidence, mortality, coverage, notification difference, TB–HIV and RR/MDR-TB"],
        ["Trends & comparisons", "Up to eight country time series plus WHO-region profile"],
        ["Age & sex", "2024 non-overlapping age bands, male/female comparison and child burden"],
        ["TB–HIV & drug resistance", "Separate rankings and a burden-versus-HIV-share scatterplot"],
        ["Forecast & data", "Exploratory three-year projection, holdout MAE, downloads and data dictionary"],
    ]
    add_table(doc, ["Dashboard tab", "Client question answered"], tab_rows, widths=[1.8, 5.0], font_size=7.6)
    add_heading(doc, "Recommended user sequence", 2)
    add_bullets(doc, [
        "Choose Global, WHO region or Country in the sidebar, then set the trend and map years.",
        "Start with the executive overview to establish scale and direction before opening rankings.",
        "Use the map and comparison tabs to identify priority geographies; use age/sex and TB–HIV/RR pages to interpret who and what is driving risk.",
        "Download the filtered data only after confirming the indicator definition and unit in the final tab.",
    ], size=8.5, space_after=1)

    # Page 4
    add_page_break(doc)
    add_kicker(doc, "03 · Reading the outputs")
    add_heading(doc, "What the dashboard shows", 1)
    add_picture(doc, CH / "global_burden_notifications.png", 6.9, "Figure 2. Global estimated burden, notifications and the resulting notification difference, 2000–2024.")
    finding_rows = [
        ["Burden", "10.55M estimated incident cases in the reproducible 2024 snapshot", "High absolute burden remains concentrated in a small group of countries."],
        ["Coverage", "8.33M notified; calculated global coverage 78.9%", "Detection and treatment access improved, but coverage remained incomplete."],
        ["Mortality", "1.20M estimated deaths; 14.7 per 100,000", "Mortality fell faster than incidence, indicating progress alongside persistent burden."],
        ["COVID-19 shock", "Notifications fell 18.1% from 2019 to 2020", "Service disruption can widen the apparent gap even when estimated incidence changes little."],
        ["Demographics", "2024 male:female estimated incidence ratio ≈ 1.50", "Adult men represent a larger share and may require tailored case-finding strategies."],
    ]
    add_table(doc, ["Signal", "Project result", "Decision meaning"], finding_rows, widths=[1.05, 2.3, 3.45], font_size=7.55, header_fill=TEAL)
    add_callout(
        doc,
        "Published WHO context",
        "WHO’s 2025 report gives rounded 2024 global headline estimates of 10.7 million incident cases and 1.23 million deaths. Small differences from this dashboard’s country-level aggregate are expected because WHO updates estimates, reports uncertainty intervals and publishes rounded totals.",
        fill="EAF5F7",
        accent=TEAL,
        font_size=8.1,
    )

    # Page 5
    add_page_break(doc)
    add_kicker(doc, "04 · Action and governance")
    add_heading(doc, "Turning dashboard signals into decisions", 1)
    action_rows = [
        ["Large absolute notification difference", "Audit subnational case finding, diagnostic access and reporting completeness; prioritize high-volume geographies."],
        ["Low calculated coverage", "Strengthen screening, rapid diagnosis, referral and linkage to treatment; investigate denominator uncertainty."],
        ["High TB–HIV share", "Integrate TB symptom screening, HIV testing, antiretroviral therapy and preventive treatment pathways."],
        ["High RR/MDR-TB burden", "Expand rapid drug-susceptibility testing, treatment access, adherence support and infection control."],
        ["Male/adult age concentration", "Adapt outreach hours, workplace/community screening and risk communication to underserved groups."],
    ]
    add_table(doc, ["Dashboard trigger", "Suggested health-system response"], action_rows, widths=[2.25, 4.55], font_size=7.7)

    add_heading(doc, "Limitations", 2)
    add_bullets(doc, [
        "Ecological country-level data cannot establish individual risk or causality.",
        "Point estimates in the analytical mirror do not include the full WHO uncertainty intervals.",
        "Age-sex analysis is limited to 2024; treatment-cohort outcomes are outside this version.",
        "Coverage can exceed 100% where estimated burden and reported counts differ because of uncertainty, timing or revisions.",
        "The forecast is a linear extrapolation for exploration only, not a WHO, clinical or causal prediction.",
    ], size=7.9, space_after=1)

    add_heading(doc, "Deployment and troubleshooting", 2)
    deploy_rows = [
        ["Local run", "pip install -r requirements.txt → streamlit run app.py"],
        ["Publish", "Upload the folder to GitHub; deploy app.py in Streamlit Community Cloud; paste the final URL into the submission."],
        ["Password", "Add APP_PASSWORD to Streamlit Secrets only when a protected landing page is required."],
        ["Data refresh", "Replace raw snapshots, run prepare_data.py and validate_data.py, then test every dashboard tab."],
    ]
    add_table(doc, ["Task", "Instruction"], deploy_rows, widths=[1.25, 5.55], font_size=7.4, header_fill=NAVY)

    add_heading(doc, "References", 2)
    refs = [
        "World Health Organization. (2025). Global tuberculosis report 2025. https://www.who.int/teams/global-programme-on-tuberculosis-and-lung-health/tb-reports/global-tuberculosis-report-2025",
        "World Health Organization. (2026). Tuberculosis data. https://www.who.int/teams/global-programme-on-tuberculosis-and-lung-health/data",
        "World Health Organization. (2026). Tuberculosis fact sheet. https://www.who.int/news-room/fact-sheets/detail/tuberculosis",
        "GTB-TME. (2025). gtbreport2025: Code and data used to develop the WHO Global tuberculosis report 2025. https://github.com/GTB-TME/gtbreport2025",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="Small")
        p.paragraph_format.left_indent = Inches(.18)
        p.paragraph_format.first_line_indent = Inches(-.18)
        p.add_run(ref)

    out = DOCS / "TB_Dashboard_Manual.docx"
    doc.save(out)
    return out


# ----------------------- REPORT -----------------------
def build_report():
    doc = setup_document(
        "Global Tuberculosis Burden and Treatment Gaps - Project Report",
        "Individual project report",
    )
    add_cover(
        doc,
        "Global Tuberculosis Burden & Treatment Gaps",
        "Trends, demographic disparities and geographic patterns",
        "Individual project report",
    )

    # Page 2 - TOC and executive summary
    add_page_break(doc)
    add_kicker(doc, "Report roadmap")
    add_heading(doc, "Table of contents", 1)
    toc_rows = [
        ["Introduction", "3"],
        ["Objective", "4"],
        ["Methodology", "4"],
        ["Results", "5"],
        ["Recommendations", "8"],
        ["References", "9"],
    ]
    add_table(doc, ["Section", "Page"], toc_rows, widths=[5.9, .8], font_size=8.5)
    add_heading(doc, "Executive summary", 1)
    add_body(doc,
        "Tuberculosis (TB) is preventable and curable, yet it remains a major global public-health problem. This project developed a published-ready Streamlit dashboard to translate WHO-derived country-level estimates and notifications into decision-oriented views of burden, trends, demographic distribution and service-coverage gaps. The core analytical file contains 5,347 country-year observations for 217 countries and territories from 2000 to 2024, supplemented with a 2024 age-sex cross-section."
    )
    add_body(doc,
        "In the reproducible 2024 snapshot, an estimated 10.55 million people developed TB, 8.33 million new and relapse cases were notified, and calculated diagnosis-and-treatment coverage was 78.9%. The resulting estimated notification difference was 2.22 million. Global incidence and mortality rates were 31.2% and 69.3% lower than in 2000, respectively, while a sharp notification fall in 2020 demonstrated the vulnerability of TB services to system-wide disruption."
    )
    add_body(doc,
        "The dashboard recommends prioritizing high-volume case-finding gaps, integrating TB and HIV pathways where co-epidemic burden is high, strengthening rapid resistance testing, and tailoring outreach to demographic groups with disproportionate burden. Findings are ecological and should guide investigation and programme prioritization rather than individual clinical decisions."
    )
    add_callout(doc, "Client value", "The tool converts a complex annual surveillance dataset into an accessible sequence: establish scale → identify priority geographies → interpret demographic and clinical dimensions → download the evidence behind the decision.", fill="EAF5F7", accent=MINT, font_size=8.8)

    # Page 3 Introduction
    add_page_break(doc)
    add_kicker(doc, "01 · Background")
    add_heading(doc, "Introduction", 1)
    add_body(doc,
        "Tuberculosis is an infectious disease caused by Mycobacterium tuberculosis. It is transmitted through the air and most commonly affects the lungs, although other organs can be involved. The disease is preventable and usually curable when people are diagnosed promptly, receive an effective drug regimen and complete treatment. Population-level TB burden therefore reflects not only transmission and underlying vulnerability, but also the reach and performance of health systems."
    )
    add_body(doc,
        "The World Health Organization (WHO) estimated that 10.7 million people developed TB in 2024 and 1.23 million died from the disease. Men accounted for 54% of incident cases, women for 35% and children for 11%. WHO also estimated approximately 619,000 incident TB cases among people living with HIV and around 390,000 cases of rifampicin-resistant or multidrug-resistant TB (RR/MDR-TB). These dimensions matter because HIV weakens immunity and increases the risk that TB infection progresses to active disease, whereas RR/MDR-TB limits the effectiveness of standard medicines and creates a more complex treatment pathway."
    )
    add_heading(doc, "The treatment-gap problem", 2)
    add_body(doc,
        "A health system cannot treat a person it does not detect or link to care. For this project, the central service-gap indicator is the difference between estimated incident TB cases and notified new and relapse cases. The ratio of notified cases to estimated incident cases is presented as calculated diagnosis-and-treatment coverage. This measure is useful for identifying potential under-detection, access barriers or reporting weaknesses, but it is not a direct count of untreated individuals because the numerator and denominator are generated through different systems and both may be revised."
    )
    add_heading(doc, "Significance for healthcare analytics", 2)
    add_body(doc,
        "TB data combine counts, population-standardized rates, time trends, demographic groups, geography and clinically important subtypes. This makes TB suitable for a healthcare analytics dashboard that must move beyond descriptive charts to support prioritization. A decision-maker needs to distinguish large absolute burden from high population risk, identify where coverage is incomplete, and understand whether HIV or drug resistance changes the required response."
    )
    add_callout(doc, "Project positioning", "The analysis is a population-health decision-support exercise. It does not diagnose TB, prescribe treatment or estimate individual risk.", fill="FFF6ED", accent=CORAL, font_size=8.8)

    # Page 4 objectives methodology
    add_page_break(doc)
    add_kicker(doc, "02 · Study design")
    add_heading(doc, "Objective", 1)
    add_body(doc, "The overall objective was to build a reproducible Streamlit dashboard that enables a healthcare decision-maker to understand the magnitude, distribution and trajectory of global TB burden and identify potential diagnosis-and-treatment coverage gaps.")
    add_heading(doc, "Specific objectives", 2)
    add_bullets(doc, [
        "Quantify and visualize estimated TB incidence, mortality, notifications and calculated coverage from 2000 to 2024.",
        "Compare countries and WHO regions using both absolute counts and population-standardized rates.",
        "Describe the 2024 distribution of estimated incident TB by age and sex.",
        "Identify countries with substantial HIV-associated TB and RR/MDR-TB burden.",
        "Provide filters, downloadable data and an optional exploratory three-year trend projection.",
    ], size=8.8, space_after=1)

    add_heading(doc, "Methodology", 1)
    add_heading(doc, "Data design and sources", 2)
    add_body(doc,
        "The project used quantitative secondary data from the WHO Global Tuberculosis Programme. Because the direct WHO CSV endpoint was not retrievable in the development runtime, WHO-derived indicator files were acquired from a public DDF mirror and checked against WHO’s current report pages, data documentation and official 2025 report repository. A separate WHO-derived snapshot was used only to attach WHO-region labels. Raw files were retained without modification, and a manifest recorded the source URL, retrieval date, file dimensions, year range and SHA-256 hash."
    )
    add_heading(doc, "Cleaning and integration", 2)
    add_body(doc,
        "Python and pandas were used to standardize country identifiers, country names, ISO codes, years and numeric values. Fifteen indicator files were merged into one country-year table. Global and regional totals were recalculated by summing counts and using population-weighted rates. Derived variables included calculated coverage, coverage gap, notification difference, case-fatality ratio, HIV-associated share and RR/MDR-TB share. Non-overlapping age bands were selected for demographic totals to prevent double-counting."
    )
    add_heading(doc, "Validation and visualization", 2)
    add_body(doc,
        "Automated validation checked the uniqueness of country-year records, year range, non-negative burden measures, ISO-code completeness and availability of male/female age-sex estimates. Streamlit and Plotly were used for the interactive dashboard; static report figures were produced in matplotlib. The optional forecast fits a linear trend to recent incidence-rate observations and reports mean absolute error (MAE) on a historical holdout period."
    )

    # Page 5 Results global
    add_page_break(doc)
    add_kicker(doc, "03 · Findings")
    add_heading(doc, "Results", 1)
    add_heading(doc, "Global burden and service coverage", 2)
    add_picture(doc, CH / "global_burden_notifications.png", 6.5, "Figure 1. Estimated incident TB, notified new and relapse cases, and the resulting notification difference, 2000–2024.")
    add_body(doc,
        "The reproducible 2024 country-level snapshot contained 10.55 million estimated incident TB cases and 8.33 million notified new and relapse cases. Calculated diagnosis-and-treatment coverage was 78.9%, leaving an estimated notification difference of 2.22 million. Although this difference was much smaller than the 7.82 million observed in the 2000 snapshot, it remained large enough to warrant targeted investigation of detection, linkage and reporting performance."
    )
    add_body(doc,
        "Notifications more than doubled between 2000 and 2024, while calculated coverage rose from 32.4% to 78.9%. In 2020, notifications fell by 18.1% and calculated coverage by 17.0%, while estimated incidence changed much less. Across 2000–2024, incidence and mortality rates declined by 31.2% and 69.3%; incidence fell only 12.3% from the 2015 End TB Strategy baseline, remaining far from the milestone."
    )
    add_picture(doc, CH / "global_rates.png", 6.5, "Figure 2. Global TB incidence and mortality rates, 2000–2024.")

    # Page 6 Results geography
    add_page_break(doc)
    add_kicker(doc, "03 · Findings continued")
    add_heading(doc, "Geographic concentration and coverage gaps", 1)
    add_picture(doc, CH / "country_notification_gaps_2024.png", 6.45, "Figure 3. Countries with the largest absolute estimated notification differences in 2024.")
    add_body(doc,
        "Eight countries dominated the absolute burden in 2024. The largest notification differences were in Indonesia (about 241,000), India (211,000), the Philippines (179,000), Pakistan (171,000) and China (165,000). India combined the largest burden with coverage above 92%, while Myanmar was below 44%. Regionally, Africa had the highest incidence rate; South-East Asia had the largest burden and highest coverage; the Western Pacific had the largest absolute notification difference. Prioritization should therefore combine risk, volume and service reach."
    )
    add_picture(doc, CH / "who_region_profile_2024.png", 6.4, "Figure 4. WHO-region profile of incidence, calculated coverage and estimated case volume in 2024.")

    # Page 7 Results demographics/clinical dimensions
    add_page_break(doc)
    add_kicker(doc, "03 · Findings continued")
    add_heading(doc, "Demographic and clinically important dimensions", 1)
    add_picture(doc, CH / "age_sex_pyramid_2024.png", 6.9, "Figure 5. Estimated incident TB by non-overlapping age band and sex, 2024.")
    add_body(doc,
        "The 2024 age-sex cross-section estimated approximately 6.34 million male and 4.22 million female incident cases, a male-to-female ratio of about 1.50. The largest adult burdens were concentrated in ages 25–54, with a pronounced male excess in ages 35–64. Children aged 0–14 accounted for approximately 1.17 million estimated cases in the non-overlapping bands used by the dashboard. These patterns support differentiated case finding rather than a single universal outreach strategy."
    )
    add_picture(doc, CH / "hiv_rr_burden_2024.png", 6.9, "Figure 6. Countries with the largest HIV-associated TB and RR/MDR-TB burdens in 2024.")
    add_body(doc,
        "HIV-associated TB was concentrated in sub-Saharan Africa. South Africa alone had an estimated 134,000 HIV-associated incident cases, representing more than half of its total estimated TB burden. RR/MDR-TB followed a different geography: India had the largest estimated count, while the European Region had the highest regional RR/MDR-TB share. This distinction matters operationally because TB–HIV requires integrated prevention, testing and antiretroviral pathways, whereas resistance requires rapid drug-susceptibility testing and access to effective regimens."
    )

    # Page 8 Recommendations / limitations
    add_page_break(doc)
    add_kicker(doc, "04 · Decision response")
    add_heading(doc, "Recommendations", 1)
    rec_rows = [
        ["1. Prioritize the gap by volume and rate", "Use the map and ranking together. High-volume countries need scaled case finding; high-rate settings need intensive risk reduction even when population size is smaller."],
        ["2. Protect continuity of essential TB services", "The 2020 notification shock supports contingency plans for diagnostics, medication supply, community follow-up and reporting during emergencies."],
        ["3. Integrate TB and HIV pathways", "In high-co-burden settings, combine TB screening, HIV testing, antiretroviral treatment, preventive therapy and patient navigation."],
        ["4. Expand rapid resistance detection", "Prioritize molecular testing, drug-susceptibility testing, treatment enrolment and adherence support where RR/MDR-TB burden or share is high."],
        ["5. Tailor outreach by demographic pattern", "Use workplace, community and flexible-hour screening for adult men while maintaining child contact investigation and preventive services."],
        ["6. Strengthen data quality governance", "Investigate values above 100% coverage, review subnational completeness, document revisions and refresh the dashboard with each new WHO release."],
    ]
    add_table(doc, ["Recommendation", "Rationale and action"], rec_rows, widths=[2.25, 4.55], font_size=7.7)

    add_heading(doc, "Limitations", 1)
    add_bullets(doc, [
        "WHO estimates have uncertainty intervals; this analytical mirror contains point estimates only.",
        "Country aggregates can differ slightly from rounded figures in the published report because of updates, scope and rounding.",
        "Notifications are programme counts, while incidence is modelled; their difference is not a verified count of untreated people.",
        "Treatment-cohort success, loss to follow-up and preventive-treatment uptake are outside the current integrated dataset.",
        "Age-sex estimates are available only for 2024 in this package, limiting demographic trend analysis.",
        "The exploratory projection assumes recent linear trend continuation and cannot anticipate policy, conflict, migration, diagnostic change or outbreaks.",
    ], size=8.2, space_after=1)
    add_heading(doc, "Conclusion", 2)
    add_body(doc,
        "The dashboard demonstrates how transparent definitions, reproducible data preparation and interactive visualization can convert global TB surveillance into practical decision support. The central message is not simply that TB burden remains high; it is that the appropriate response depends on the combination of burden volume, population rate, service coverage, demographic distribution, HIV co-burden and drug resistance."
    )

    # Page 9 references
    add_page_break(doc)
    add_kicker(doc, "05 · Evidence base")
    add_heading(doc, "References", 1)
    refs = [
        "GTB-TME. (2025). gtbreport2025: Code and data used to develop the World Health Organization’s Global tuberculosis report 2025. GitHub. https://github.com/GTB-TME/gtbreport2025",
        "Open Numbers. (2026). DDF dataset: WHO TB burden estimates. https://github.com/open-numbers/ddf--who--tb_burden_estimates",
        "World Health Organization. (2015). The End TB Strategy. Geneva: WHO. https://www.who.int/teams/global-programme-on-tuberculosis-and-lung-health/the-end-tb-strategy",
        "World Health Organization. (2025). Global tuberculosis report 2025. Geneva: WHO. https://www.who.int/teams/global-programme-on-tuberculosis-and-lung-health/tb-reports/global-tuberculosis-report-2025",
        "World Health Organization. (2025). Top findings and messages in the 2025 report. https://cdn.who.int/media/docs/default-source/global-tuberculosis-report-2025/top-findings-and-messages-in-the-2025-report_english.pdf",
        "World Health Organization. (2026). Tuberculosis data. https://www.who.int/teams/global-programme-on-tuberculosis-and-lung-health/data",
        "World Health Organization. (2026, March 24). Tuberculosis fact sheet. https://www.who.int/news-room/fact-sheets/detail/tuberculosis",
        "Wickham, H., Çetinkaya-Rundel, M., & Grolemund, G. (2023). R for Data Science (2nd ed.). O’Reilly Media. [Referenced for reproducible data workflow principles.]",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(.25)
        p.paragraph_format.first_line_indent = Inches(-.25)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(ref)
        set_run_font(r, 8.5, False, NAVY)

    add_heading(doc, "Project data files", 2)
    add_body(doc,
        "The submitted ZIP includes raw snapshots, processed country-year and age-sex files, a data dictionary, a source manifest with hashes, preparation and validation scripts, and the complete Streamlit source code. The final published dashboard URL should be inserted after deployment: [PASTE STREAMLIT URL HERE].",
        style="Small",
    )
    add_callout(doc, "Version note", "Prepared 23 June 2026 using data retrieved on the same date. WHO updates TB estimates annually; refresh the source snapshots before reusing the dashboard for future reporting.", fill="EAF5F7", accent=TEAL, font_size=8.5)

    out = DOCS / "TB_Individual_Project_Report.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_manual())
    print(build_report())
